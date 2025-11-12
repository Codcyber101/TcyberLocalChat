"""
DocumentService for file upload, processing, and management
"""

import os
import shutil
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

import aiofiles
from fastapi import Depends, UploadFile
from sqlalchemy.orm import Session

# Feature-detection flags for optional PDF libraries
_HAS_PYMUPDF = False
_HAS_PYPDF = False
try:
    # pypdf is the maintained successor to PyPDF2
    from pypdf import PdfReader

    _HAS_PYPDF = True
except Exception:
    # Fallback to PyMuPDF or PyPDF2 if pypdf isn't installed
    try:
        # Prefer PyMuPDF for robust, fast PDF text extraction
        import fitz  # PyMuPDF

        _HAS_PYMUPDF = True
        PdfReader = None
        _HAS_PYPDF = False
    except Exception:
        _HAS_PYMUPDF = False
        # Try pypdf again (some environments may have it under different discovery paths)
        try:
            from pypdf import PdfReader

            _HAS_PYPDF = True
        except Exception:
            # Try PyPDF2 as a last resort; if it's missing, gracefully set PdfReader=None
            try:
                from PyPDF2 import PdfReader

                _HAS_PYPDF = True
            except Exception:
                PdfReader = None
                _HAS_PYPDF = False
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import cv2
except Exception:
    cv2 = None

try:
    import numpy as np
except Exception:
    np = None

from ..database import get_db
from ..models.document import Document


class DocumentService:
    """Service for managing document upload, processing, and storage"""

    # Supported MIME types
    SUPPORTED_TEXT_TYPES = ["text/plain", "text/markdown", "application/pdf"]
    SUPPORTED_IMAGE_TYPES = ["image/jpeg", "image/png", "image/gif", "image/webp"]
    SUPPORTED_AUDIO_TYPES = ["audio/mpeg", "audio/wav", "audio/ogg", "audio/mp4"]

    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_IMAGE_SIZE = 10 * 1024 * 1024  # 10MB for images
    MAX_AUDIO_SIZE = 10 * 1024 * 1024  # 10MB for audio

    def __init__(self, db: Session, upload_dir: str = "uploads"):
        self.db = db
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)

    def validate_file(self, file: UploadFile) -> tuple[bool, str | None]:
        """Validate uploaded file type and size"""
        # Check file size
        if hasattr(file, "size") and file.size:
            mime = file.content_type or "application/octet-stream"
            max_size = self._get_max_size_for_mime_type(mime)
            if file.size > max_size:
                return (
                    False,
                    f"File too large. Maximum size: {max_size // (1024 * 1024)}MB",
                )

        # Check MIME type
        mime_type = file.content_type or ""
        if mime_type not in self._get_supported_types():
            return False, f"Unsupported file type: {mime_type}"

        return True, None

    def _get_max_size_for_mime_type(self, mime_type: str) -> int:
        """Get maximum file size for MIME type"""
        if not mime_type:
            return self.MAX_FILE_SIZE

        if isinstance(mime_type, str) and mime_type.startswith("image/"):
            return self.MAX_IMAGE_SIZE
        elif isinstance(mime_type, str) and mime_type.startswith("audio/"):
            return self.MAX_AUDIO_SIZE
        else:
            return self.MAX_FILE_SIZE

    def _get_supported_types(self) -> list[str]:
        """Get all supported MIME types"""
        return (
            self.SUPPORTED_TEXT_TYPES
            + self.SUPPORTED_IMAGE_TYPES
            + self.SUPPORTED_AUDIO_TYPES
        )

    def save_uploaded_file(self, file: UploadFile) -> str:
        """Save uploaded file and return file path"""
        file_id = str(uuid4())
        file_extension = Path(file.filename).suffix if file.filename else ""

        # Create subdirectory structure for organization
        subdir = self.upload_dir / file_id[:2] / file_id[2:4]
        subdir.mkdir(parents=True, exist_ok=True)

        file_path = subdir / f"{file_id}{file_extension}"

        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        return str(file_path)

    def create_document(
        self,
        file_name: str,
        file_content: str,
        status: str = "processing",
        mime_type: str = None,
    ) -> Document:
        """Create a new document, save its content, and create a database record."""
        # For simplicity, let's save the content to a temporary file
        # In a real application, you might store content in a dedicated storage
        file_id = str(uuid4())
        file_path = self.upload_dir / f"{file_id}_{file_name}"
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(file_content)

        # Auto-detect MIME type from extension if not provided
        if mime_type is None:
            filename_lower = file_name.lower()
            if filename_lower.endswith((".md", ".markdown")):
                mime_type = "text/markdown"
            elif filename_lower.endswith(".txt"):
                mime_type = "text/plain"
            elif filename_lower.endswith(".pdf"):
                mime_type = "application/pdf"
            else:
                mime_type = "text/plain"  # Default for text files

        document = Document(
            filename=file_name,
            mime_type=mime_type,
            path=str(file_path),
            size=len(file_content.encode("utf-8")),
            content=file_content,
            status=status,
        )

        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)

        return document

    def create_document_record(
        self, filename: str, file_path: str, size: int, mime_type: str
    ) -> Document:
        """Create database record for uploaded document"""
        document = Document(
            filename=filename,
            path=file_path,
            size=size,
            mime_type=mime_type,
            status="processing",
        )

        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)

        return document

    def get_document(self, document_id: str) -> Document | None:
        """Get document by ID"""
        try:
            doc_uuid = UUID(document_id)
        except ValueError:
            return None
        # Ensure we return ORM instance (avoid returning Column proxies)
        return self.db.query(Document).filter(Document.id == doc_uuid).first()

    def get_all_documents(self) -> list[Document]:
        """Get all documents"""
        return self.db.query(Document).order_by(Document.uploaded_at.desc()).all()

    def get_documents(
        self, status: str | None = None, limit: int = 50
    ) -> list[Document]:
        """Get documents with optional status filter"""
        query = self.db.query(Document)
        if status:
            query = query.filter(Document.status == status)
        return query.order_by(Document.uploaded_at.desc()).limit(limit).all()

    def update_document_status(self, document_id: str, status: str) -> Document | None:
        """Update document processing status"""
        document = self.get_document(document_id)
        if document:
            document.update_status(status)
            self.db.commit()
            self.db.refresh(document)
        return document

    def update_document_filename(
        self, document_id: str, new_filename: str
    ) -> Document | None:
        """Update document filename"""
        document = self.get_document(document_id)
        if document:
            document.filename = new_filename
            self.db.commit()
            self.db.refresh(document)
        return document

    def delete_document(self, document_id: str) -> bool:
        """Delete document and its file"""
        document = self.get_document(document_id)
        if document:
            # Delete file from disk
            try:
                if os.path.exists(str(document.path)):
                    os.remove(str(document.path))
            except OSError:
                pass  # File may not exist or can't be deleted

            # Delete database record
            self.db.delete(document)
            self.db.commit()
            return True
        return False

    async def process_document_async(self, document_id: str) -> None:
        """Async processing pipeline for documents"""
        document = self.get_document(document_id)
        if not document:
            return

        try:
            # Update status to processing
            document.update_status("processing")

            # Extract text content based on file type
            # Coerce ORM attributes into local Python values for safe checks
            is_text = bool(getattr(document, "is_text_document", False))
            is_image = bool(getattr(document, "is_image", False))
            is_audio = bool(getattr(document, "is_audio", False))

            if is_text:
                await self._extract_text_content(document)
            elif is_image:
                await self._process_image(document)
            elif is_audio:
                await self._transcribe_audio(document)

            # Generate chunks for RAG
            content_val = getattr(document, "content", None)
            if content_val:
                await self._generate_chunks(document)

            # Update status to ready and then completed
            document.update_status("ready")
            # Some consumers expect a final 'completed' state
            document.update_status("completed")

        except Exception:
            document.update_status("error")
            # Log error would go here

        self.db.commit()

    async def _extract_text_content(self, document: Document) -> None:
        """Extract text from text-based documents"""
        try:
            content = ""
            mime = getattr(document, "mime_type", "") or ""

            # Check if content is already in memory (from direct upload)
            existing_content = getattr(document, "content", None)
            file_path = str(getattr(document, "path", ""))

            if mime == "application/pdf":
                if existing_content:
                    # PDF content already extracted, use it
                    content = existing_content
                else:
                    content = await self._extract_pdf_content(file_path)
            elif mime == "text/markdown":
                if existing_content:
                    # Markdown content already in memory, process it directly
                    content = self._process_markdown_content(existing_content)
                else:
                    # Read from file and process
                    content = await self._extract_markdown_content(file_path)
            elif mime == "text/plain":
                if existing_content:
                    # Plain text already in memory, use it
                    content = existing_content
                else:
                    content = await self._extract_text_file_content(file_path)
            elif (
                mime
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                content = await self._extract_docx_content(file_path)
            else:
                content = f"Unsupported text format: {mime}"

            document.content = content

        except Exception as e:
            document.content = f"Error extracting content: {str(e)}"
            raise

    async def _process_image(self, document: Document) -> None:
        """Process image files (OCR, analysis)"""
        # Placeholder for image processing
        document.content = f"Image content from {document.filename}"

    async def _transcribe_audio(self, document: Document) -> None:
        """Transcribe audio files"""
        # Placeholder for audio transcription
        document.transcription = f"Transcription of {document.filename}"

    async def _generate_chunks(self, document: Document) -> None:
        """Generate text chunks for RAG and create vector embeddings"""
        content_val = getattr(document, "content", None)
        if not content_val:
            print(f"DEBUG: _generate_chunks - No content for document {document.id}")
            return

        # Import RAG service
        from .rag_service import get_rag_service

        rag_service = get_rag_service()

        # Use RAG service to add document with intelligent chunking
        # This will automatically create chunks and embeddings
        # Prepare safe metadata values
        doc_id = str(document.id)
        full_text = str(getattr(document, "content", ""))
        filename = str(getattr(document, "filename", ""))
        mime = str(getattr(document, "mime_type", ""))
        uploaded_at = getattr(document, "uploaded_at", None)
        size_val = getattr(document, "size", None)
        try:
            file_size = int(size_val) if size_val is not None else 0
        except Exception:
            file_size = 0

        print(
            f"DEBUG: _generate_chunks - Adding document_id={doc_id}, filename={filename}, content_length={len(full_text)}"
        )
        success = await rag_service.add_document_with_chunking(
            document_id=doc_id,
            full_text=full_text,
            metadata={
                "filename": filename,
                "mime_type": mime,
                "upload_date": uploaded_at.isoformat() if uploaded_at else None,
                "file_size": file_size,
            },
        )
        print(
            f"DEBUG: _generate_chunks - add_document_with_chunking returned success={success} for document_id={doc_id}"
        )

        if success:
            # Mark document as having embeddings
            document.has_embeddings = True
            # Store chunk count (approximate)
            # Note: In a real implementation, you'd get this from the RAG service
            document.chunks = [str(getattr(document, "content", ""))]  # Placeholder
        else:
            # Fallback to simple chunking if RAG service fails
            await self._fallback_chunking(document)

    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text content from PDF file"""
        try:
            # Use PyMuPDF if available for more reliable extraction
            if _HAS_PYMUPDF:
                try:
                    doc = fitz.open(file_path)
                    text_parts = []
                    for page in doc:
                        try:
                            text_parts.append(page.get_text("text") or "")
                        except Exception:
                            # best-effort per-page
                            try:
                                text_parts.append(page.get_text() or "")
                            except Exception:
                                continue
                    return "\n".join(text_parts).strip()
                except Exception:
                    # Fall through to pypdf fallback if something unexpected happens
                    pass

            # Fallback to pypdf / PyPDF2 reader
            if PdfReader is not None:
                reader = PdfReader(file_path)
                text = ""
                # pypdf/PyPDF2 expose pages differently across versions
                try:
                    for page in getattr(reader, "pages", []):
                        txt = None
                        try:
                            txt = page.extract_text()
                        except Exception:
                            try:
                                txt = page.get_text()
                            except Exception:
                                txt = None
                        if txt:
                            text += txt + "\n"
                except Exception:
                    # Older PyPDF2 may require iterating differently
                    try:
                        for i in range(len(reader.pages)):
                            p = reader.pages[i]
                            txt = (
                                p.extract_text() if hasattr(p, "extract_text") else None
                            )
                            if txt:
                                text += txt + "\n"
                    except Exception:
                        return "Error extracting PDF content: unsupported pdf reader format"

                return text.strip()

            return ""
        except Exception as e:
            return f"Error extracting PDF content: {str(e)}"

    async def _extract_text_file_content(self, file_path: str) -> str:
        """Extract text content from plain text file"""
        try:
            async with aiofiles.open(file_path, encoding="utf-8") as f:
                return await f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, encoding="latin-1") as f:
                    return await f.read()
            except Exception as e:
                return f"Error reading text file: {str(e)}"
        except Exception as e:
            return f"Error reading text file: {str(e)}"

    def _process_markdown_content(self, content: str) -> str:
        """Process markdown content (frontmatter parsing) - works with content already in memory"""
        # Parse frontmatter if present (YAML frontmatter between --- markers)
        frontmatter = {}
        body = content

        if content.startswith("---"):
            try:
                # Split frontmatter from body
                parts = content.split("---", 2)
                if len(parts) >= 3:
                    frontmatter_text = parts[1].strip()
                    body = parts[2].strip()

                    # Simple YAML frontmatter parsing (basic key-value pairs)
                    # For more complex YAML, we'd need PyYAML, but keeping it simple for now
                    for line in frontmatter_text.split("\n"):
                        if ":" in line:
                            key, value = line.split(":", 1)
                            frontmatter[key.strip()] = (
                                value.strip().strip('"').strip("'")
                            )
            except Exception:
                # If frontmatter parsing fails, use entire content as body
                pass

        # Return content with frontmatter metadata preserved in a comment if present
        # This ensures frontmatter data is available for RAG but doesn't interfere with markdown structure
        if frontmatter:
            frontmatter_comment = f"<!-- Frontmatter: {str(frontmatter)} -->\n\n"
            return frontmatter_comment + body

        return body

    async def _extract_markdown_content(self, file_path: str) -> str:
        """Extract and process markdown content from file, including frontmatter parsing"""
        try:
            async with aiofiles.open(file_path, encoding="utf-8") as f:
                content = await f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, encoding="latin-1") as f:
                    content = await f.read()
            except Exception as e:
                return f"Error reading markdown file: {str(e)}"
        except Exception as e:
            return f"Error reading markdown file: {str(e)}"

        # Use the shared processing method
        return self._process_markdown_content(content)

    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text content from Word document"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            return f"Error extracting DOCX content: {str(e)}"

    async def _extract_image_text(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            # Read image
            image = cv2.imread(file_path)

            if image is None:
                return "Error: Image not found or could not be read"

            # Convert to grayscale for better OCR
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

            # Apply threshold to get better contrast
            _, threshold = cv2.threshold(
                gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
            )

            # OCR the image
            text = pytesseract.image_to_string(threshold)

            return text.strip()
        except Exception as e:
            return f"Error extracting image text: {str(e)}"

    async def _fallback_chunking(self, document: Document) -> None:
        """Fallback chunking when RAG service is unavailable"""
        content_val = getattr(document, "content", "") or ""
        if content_val:
            # Simple sentence-based chunking
            sentences = content_val.split(".")
            chunks = []
            current_chunk = ""

            for sentence in sentences:
                if len(current_chunk + sentence) > 1000:  # Chunk size limit
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence
                else:
                    current_chunk += sentence + "."

            if current_chunk:
                chunks.append(current_chunk.strip())

            document.chunks = chunks

    async def render_as_html(
        self,
        document: Document,
        include_metadata: bool = True,
        highlight_terms: list[str] | None = None,
        options: dict[str, Any] | None = None,
        page: int | None = None,
    ) -> str:
        """Render document as HTML"""
        content = getattr(document, "content", "") or ""

        # Apply highlighting if requested
        if highlight_terms:
            for term in highlight_terms:
                content = content.replace(
                    term, f'<mark class="highlight">{term}</mark>'
                )

        # Pagination (simple implementation)
        if page is not None:
            sentences = content.split(".")
            page_size = 10
            start_idx = page * page_size
            end_idx = start_idx + page_size
            page_sentences = sentences[start_idx:end_idx]
            content = ". ".join(page_sentences)

        html = '<div class="document-content">'
        if include_metadata:
            filename = getattr(document, "filename", "")
            size_val = getattr(document, "size", 0)
            mime = getattr(document, "mime_type", "")
            html += '<div class="document-metadata">'
            html += f"<h3>{filename}</h3>"
            html += (
                f'<p class="document-info">Size: {size_val} bytes | Type: {mime}</p>'
            )
            html += "</div>"

        html += f'<div class="document-body">{content}</div>'
        html += "</div>"

        return html

    async def render_as_json(
        self, document: Document, include_metadata: bool = True, page: int | None = None
    ) -> str:
        """Render document as JSON"""
        import json

        data = {
            "id": str(getattr(document, "id", "")),
            "filename": getattr(document, "filename", ""),
            "content": getattr(document, "content", ""),
            "mimeType": getattr(document, "mime_type", ""),
            "size": getattr(document, "size", 0),
            "status": getattr(document, "status", ""),
            "uploadedAt": getattr(document, "uploaded_at", None).isoformat()
            if getattr(document, "uploaded_at", None)
            else None,
        }

        if include_metadata:
            data.update(
                {
                    "chunks": document.chunks,
                    "hasEmbeddings": getattr(document, "has_embeddings", False),
                    "processingTime": getattr(document, "processing_time", None),
                }
            )

        return json.dumps(data, indent=2, default=str)

    async def render_as_text(
        self, document: Document, include_metadata: bool = True, page: int | None = None
    ) -> str:
        """Render document as plain text"""
        content = getattr(document, "content", "") or ""

        if include_metadata:
            header = f"Document: {getattr(document, 'filename', '')}\n"
            header += f"Size: {getattr(document, 'size', 0)} bytes\n"
            header += f"Type: {getattr(document, 'mime_type', '')}\n"
            header += f"Status: {getattr(document, 'status', '')}\n"
            header += "-" * 50 + "\n\n"
            content = header + content

        return content

    async def render_as_markdown(
        self, document: Document, include_metadata: bool = True, page: int | None = None
    ) -> str:
        """Render document as Markdown"""
        content = getattr(document, "content", "") or ""

        markdown = ""
        if include_metadata:
            markdown += f"# {getattr(document, 'filename', '')}\n\n"
            markdown += "**File Information:**\n"
            markdown += f"- Size: {getattr(document, 'size', 0)} bytes\n"
            markdown += f"- Type: {getattr(document, 'mime_type', '')}\n"
            markdown += f"- Status: {getattr(document, 'status', '')}\n"
            uploaded_at_val = getattr(document, "uploaded_at", None)
            markdown += f"- Uploaded: {uploaded_at_val.strftime('%Y-%m-%d %H:%M:%S') if uploaded_at_val else 'Unknown'}\n\n"
            markdown += "---\n\n"

        markdown += f"## Content\n\n{content}"

        return markdown

    async def generate_preview(
        self, document: Document, max_length: int = 500, include_metadata: bool = True
    ) -> dict[str, Any]:
        """Generate a preview of the document"""
        content = getattr(document, "content", "") or ""
        preview_text = content[:max_length]
        if len(content) > max_length:
            preview_text += "..."

        preview = {
            "id": str(getattr(document, "id", "")),
            "filename": getattr(document, "filename", ""),
            "preview": preview_text,
            "contentLength": len(content),
            "mimeType": getattr(document, "mime_type", ""),
            "size": getattr(document, "size", 0),
        }

        if include_metadata:
            uploaded_at_val = getattr(document, "uploaded_at", None)
            preview.update(
                {
                    "status": getattr(document, "status", ""),
                    "uploadedAt": uploaded_at_val.isoformat()
                    if uploaded_at_val
                    else None,
                    "hasEmbeddings": getattr(document, "has_embeddings", False),
                }
            )

        return preview

    async def get_page_info(self, document: Document) -> dict[str, Any]:
        """Get pagination information for the document"""
        content = getattr(document, "content", "") or ""
        # Simple word-based estimation
        words = len(content.split())
        estimated_pages = max(1, words // 250)  # Assume ~250 words per page

        return {
            "total_pages": estimated_pages,
            "estimated_total_length": len(content),
            "word_count": words,
            "has_pages": estimated_pages > 1,
        }

    async def extract_sections(
        self, document: Document, sections: list[str], format: str = "json"
    ) -> dict[str, Any]:
        """Extract specific sections from document"""
        content = document.content or ""

        extracted = {"document_id": str(document.id), "sections": {}}

        for section in sections:
            if section == "headers":
                # Simple header extraction (lines that are short and end with colon)
                headers = [
                    line.strip()
                    for line in content.split("\n")
                    if len(line.strip()) < 100 and line.strip().endswith(":")
                ]
                extracted["sections"]["headers"] = headers

            elif section == "paragraphs":
                # Split by double newlines or long lines
                paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
                extracted["sections"]["paragraphs"] = paragraphs

            elif section == "tables":
                # Placeholder for table extraction
                extracted["sections"]["tables"] = []

            elif section == "images":
                # Placeholder for image extraction
                extracted["sections"]["images"] = []

        if format == "json":
            return extracted
        else:
            return extracted

    async def search_within_document(
        self,
        document: Document,
        query: str,
        case_sensitive: bool = False,
        whole_words: bool = False,
    ) -> dict[str, Any]:
        """Search for text within a document"""
        content = document.content or ""

        if not case_sensitive:
            content = content.lower()
            query = query.lower()

        # Simple search implementation
        matches = []
        lines = content.split("\n")
        for i, line in enumerate(lines):
            if whole_words:
                # Word boundary search
                import re

                pattern = r"\b" + re.escape(query) + r"\b"
                if re.search(pattern, line):
                    matches.append(
                        {
                            "line": i + 1,
                            "text": line,
                            "start": line.find(query),
                            "end": line.find(query) + len(query),
                        }
                    )
            else:
                # Simple substring search
                start = line.find(query)
                if start != -1:
                    matches.append(
                        {
                            "line": i + 1,
                            "text": line,
                            "start": start,
                            "end": start + len(query),
                        }
                    )

        return {
            "query": query,
            "total_matches": len(matches),
            "matches": matches[:50],  # Limit results
            "case_sensitive": case_sensitive,
            "whole_words": whole_words,
        }


# Dependency injection function
def get_document_service(db: Session = Depends(get_db)) -> DocumentService:
    """Get DocumentService instance with database session"""
    return DocumentService(db)
